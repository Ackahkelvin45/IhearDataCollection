import os
import hashlib
from typing import List, Dict, Any, Tuple
from pathlib import Path
import PyPDF2
import docx
from PIL import Image
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats"""

    def __init__(self):
        from django.conf import settings

        self.config = settings.CHATBOT["DOCUMENT"]
        self.max_file_size = self.config["MAX_FILE_SIZE_MB"] * 1024 * 1024
        self.allowed_extensions = self.config["ALLOWED_EXTENSIONS"]
        self.enable_ocr = self.config["ENABLE_OCR"]

    def validate_file(self, file) -> Tuple[bool, str]:
        """
        Validate uploaded file

        Returns:
            (is_valid, error_message)
        """
        # Check file size
        if file.size > self.max_file_size:
            return (
                False,
                f"File size exceeds maximum allowed size of {self.config['MAX_FILE_SIZE_MB']}MB",
            )

        # Check file extension
        ext = Path(file.name).suffix.lower().lstrip(".")
        if ext not in self.allowed_extensions:
            return (
                False,
                f"File type .{ext} not allowed. Allowed types: {', '.join(self.allowed_extensions)}",
            )

        return (True, "")

    def extract_text_from_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from various file formats

        Returns:
            (extracted_text, metadata)
        """
        ext = Path(file_path).suffix.lower()
        metadata = {
            "file_type": ext.lstrip("."),
            "file_size": os.path.getsize(file_path),
        }

        try:
            if ext == ".pdf":
                text, pdf_metadata = self._extract_from_pdf(file_path)
                metadata.update(pdf_metadata)
            elif ext in [".docx", ".doc"]:
                text, doc_metadata = self._extract_from_docx(file_path)
                metadata.update(doc_metadata)
            elif ext in [".txt", ".md"]:
                text = self._extract_from_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")

            metadata["total_characters"] = len(text)
            return text, metadata

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text and images from PDF"""
        text = ""
        images_data = []
        metadata = {}

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["total_pages"] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    text += f"\n\n--- Page {page_num} ---\n\n{page_text}"

                    # Extract images if OCR is enabled
                    if self.enable_ocr:
                        try:
                            # Extract images from page
                            if "/XObject" in page["/Resources"]:
                                x_objects = page["/Resources"]["/XObject"].get_object()

                                for obj in x_objects:
                                    if x_objects[obj]["/Subtype"] == "/Image":
                                        images_data.append(
                                            {
                                                "page": page_num,
                                                "object_name": obj,
                                            }
                                        )
                        except Exception as e:
                            logger.warning(
                                f"Could not extract images from page {page_num}: {e}"
                            )

            metadata["images_found"] = len(images_data)
            metadata["images"] = images_data

            return text, metadata

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

    def _extract_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)

            # Extract paragraphs
            paragraphs = [
                paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
            ]
            text = "\n\n".join(paragraphs)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)

            if tables_text:
                text += "\n\n--- Tables ---\n\n" + "\n".join(tables_text)

            metadata = {
                "total_paragraphs": len(paragraphs),
                "total_tables": len(doc.tables),
                "has_images": len(doc.inline_shapes) > 0,
                "images_count": len(doc.inline_shapes),
            }

            return text, metadata

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text or markdown files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    def perform_ocr_on_image(self, image_path: str) -> str:
        """
        Perform OCR on an image (placeholder for future implementation)

        Note: This requires pytesseract or cloud OCR service
        """
        # Placeholder for OCR functionality
        # In production, you would use:
        # - pytesseract (local)
        # - Google Cloud Vision API
        # - Azure Computer Vision
        # - AWS Textract

        logger.info(f"OCR requested for {image_path} - Feature not yet implemented")
        return ""

    def calculate_content_hash(self, content: str) -> str:
        """Calculate SHA256 hash of content"""
        return hashlib.sha256(content.encode()).hexdigest()

    def chunk_text(
        self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """
        Split text into chunks with smart boundary detection

        Returns:
            List of chunk dictionaries with content and metadata
        """
        from langchain.text_splitter import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        )

        chunks = splitter.split_text(text)

        # Add metadata to each chunk
        chunk_data = []
        for idx, chunk in enumerate(chunks):
            chunk_data.append(
                {
                    "content": chunk,
                    "chunk_index": idx,
                    "chunk_size": len(chunk),
                    "total_chunks": len(chunks),
                }
            )

        return chunk_data

    def extract_metadata_from_chunk(self, chunk_text: str) -> Dict[str, Any]:
        """
        Extract metadata from chunk (section titles, page numbers, etc.)
        """
        metadata = {}

        # Extract page number if present
        import re

        page_match = re.search(r"---\s*Page\s+(\d+)\s*---", chunk_text)
        if page_match:
            metadata["page_number"] = int(page_match.group(1))

        # Try to detect section headers (lines starting with #, ##, etc. or all caps)
        lines = chunk_text.split("\n")
        for line in lines[:3]:  # Check first 3 lines
            line = line.strip()
            if line.startswith("#"):
                # Markdown header
                metadata["section_title"] = line.lstrip("#").strip()
                break
            elif line.isupper() and len(line) > 5 and len(line) < 100:
                # All caps header
                metadata["section_title"] = line
                break

        return metadata

    def process_document(
        self, file_path: str, doc_id: str, title: str
    ) -> Dict[str, Any]:
        """
        Complete document processing pipeline

        Args:
            file_path: Path to the document file
            doc_id: Document UUID
            title: Document title

        Returns:
            Processing result with chunks and metadata
        """
        from django.conf import settings

        try:
            # Extract text and metadata
            text, file_metadata = self.extract_text_from_file(file_path)

            # Calculate content hash
            content_hash = self.calculate_content_hash(text)

            # Chunk the text
            chunk_config = settings.CHATBOT["RAG"]
            chunks = self.chunk_text(
                text,
                chunk_size=chunk_config["CHUNK_SIZE"],
                chunk_overlap=chunk_config["CHUNK_OVERLAP"],
            )

            # Enhance chunk metadata
            for chunk in chunks:
                chunk["doc_id"] = doc_id
                chunk["title"] = title
                extracted_meta = self.extract_metadata_from_chunk(chunk["content"])
                chunk["metadata"] = extracted_meta

            result = {
                "content_hash": content_hash,
                "chunks": chunks,
                "total_chunks": len(chunks),
                "total_characters": len(text),
                "file_metadata": file_metadata,
            }

            logger.info(
                f"Processed document {doc_id}: {len(chunks)} chunks, {len(text)} chars"
            )
            return result

        except Exception as e:
            logger.error(f"Error processing document {doc_id}: {e}")
            raise
